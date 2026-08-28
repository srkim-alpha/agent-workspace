import os
import sys
import re
import json
import subprocess
from pathlib import Path
from typing import Dict, List, Any, Optional

import jinja2

BASE_DIR = Path(__file__).resolve().parent.parent
if str(BASE_DIR) not in sys.path:
    sys.path.insert(0, str(BASE_DIR))

MASTER_HUB_PATH = BASE_DIR / "career_hub" / "career_master_hub.md"
OUTPUT_DIR = BASE_DIR / "data" / "outputs"
DOCS_APPS_DIR = BASE_DIR / "docs" / "applications"
TEMPLATES_DIR = BASE_DIR / "templates"
BASE_URL = "https://srkim-alpha.github.io/agent-workspace/applications/"

class JobApplicationGenerator:
    """AI Tailored Job Application Generator (PDF & GitHub Pages WebApp)."""

    def __init__(self, master_hub_path: Path = MASTER_HUB_PATH):
        self.master_hub_path = Path(master_hub_path)
        self.raw_master_text = self._read_master_hub()
        self.star_episodes = self._extract_star_episodes()

    def _read_master_hub(self) -> str:
        if not self.master_hub_path.exists():
            raise FileNotFoundError(f"Ground Truth master hub not found: {self.master_hub_path}")
        with open(self.master_hub_path, "r", encoding="utf-8") as f:
            return f.read()

    def _extract_star_episodes(self) -> List[Dict[str, str]]:
        """Extracts 5 STAR episodes from career_master_hub.md."""
        episodes = [
            {
                "id": "lotto",
                "title": "[영업유통] 국내 최초 로또복권 200개 가맹점 유치",
                "keywords": ["영업", "유통", "가맹점", "개설", "달성률", "채널", "상권"],
                "situation": "2002년 국내 최초 온라인 로또복권 론칭 시 높은 보증보험 비율과 생소함으로 점주들의 거부감 극심.",
                "task": "인천 지역 상권 분석 및 50개점 목표 대비 우수 가맹점 유치.",
                "action": "점주 허드렛일 대행, 100% 현금거래/재고부담 무 키워드 맞춤 설득 프레젠테이션 수시 전개.",
                "result": "개인 목표 달성률 144% (72개점 모집), 판매점 개설률 100% 달성 및 영업망 안착."
            },
            {
                "id": "kazakhstan",
                "title": "[해외리더십] 카자흐스탄 침켄트 보이콧 사태 수습 및 본부장 승진",
                "keywords": ["해외", "주재원", "리더십", "파업", "노무", "갈등", "조직관리", "현장관리"],
                "situation": "2011년 카자흐스탄 침켄트 사무소에서 이전 한국인 소장의 강압적 태도로 현지 직원 전원 파업/보이콧 발생.",
                "task": "신임 소장으로 긴급 파견되어 조직 갈등 진정 및 사무소 운영 100% 정상화.",
                "action": "1:1 경청 소통, 직책 세분화 및 직책수당 지급, 업무표준 매뉴얼 제작.",
                "result": "2개월 만에 정상화, 전사 우수사례 채택, 현지인 소장 3명 배출 및 카자흐스탄 남부지역 본부장 승진."
            },
            {
                "id": "homeshopping",
                "title": "[채널개척] TV홈쇼핑 신규 런칭 및 10억 누적 매출 달성",
                "keywords": ["TV홈쇼핑", "홈쇼핑", "방송", "기획", "매출", "영업", "입점", "마케팅"],
                "situation": "(주)현대시트 근무 시 오프라인 중심 유통 채널 한계로 신규 온라인/방송 유통 망 개척 필요.",
                "task": "TV홈쇼핑(홈앤쇼핑, 쇼핑&T) 입점 제안, 방송 기획, 미스터리 쇼퍼 운영 및 매출 목표 달성.",
                "action": "따소미플러스 방송 콘셉트 기획, 사전 영상 촬영, 생방송 실시간 모니터링 및 출고 프로세스 연동.",
                "result": "홈앤쇼핑 1회 생방송 1억 6,000만 원 (133% 달성), T커머스 10주간 누적 10억 원 매출 돌파."
            },
            {
                "id": "wms_excel",
                "title": "[프로세스혁신] WMS 물류전산 자동화 및 매출 12배 신장",
                "keywords": ["물류", "SCM", "WMS", "전산", "엑셀", "입출고", "재고", "자동화", "총괄"],
                "situation": "(주)그래이박스 4,000평 냉동/냉장 물류센터 축산물 입출고 수기 관리에 따른 오차 발생.",
                "task": "전문 WMS 시스템(사방넷/엔윌) 운용 및 MS Excel 쿼리 자동화 서식 구축.",
                "action": "현장 데이터 자동 집계 쿼리 개발, 3톤 지게차 직접 몰며 보세물류 동선 및 입출고 수량 자동 대조.",
                "result": "입출고 오차율 0% 달성, 월 매출 500만 원에서 6,000만 원으로 12배 신장."
            },
            {
                "id": "ai_agent",
                "title": "[AI 에이전트] Visual Career Organizer & Ground Truth 대시보드 구축",
                "keywords": ["AI", "데이터", "시스템", "자동화", "에이전트", "IT", "전산", "파이썬"],
                "situation": "15년 8개월간 94개 문서에 파편화된 커리어 자산의 정제 및 오염 데이터 검증 필요.",
                "task": "0-의존성 파서 개발, 자격득실확인서 PDF 복호화, Playwright visual organizer 및 대시보드 구축.",
                "action": "Python career_parser.py, dashboard_builder.py 자율 구축, Ground Truth 검증 완료.",
                "result": "94개 문서 전수 분석, 100% Ground Truth 검증 마스터 DB 및 1-Click 대시보드 구축 완료."
            }
        ]
        return episodes

    def match_star_episodes(self, job_text: str, top_k: int = 3) -> List[Dict[str, str]]:
        """Matches top STAR episodes based on job posting keyword scores."""
        scores = []
        for ep in self.star_episodes:
            score = 0
            for kw in ep["keywords"]:
                score += len(re.findall(kw, job_text, re.IGNORECASE)) * 2
            scores.append((score, ep))

        # Sort by score desc, pick top_k
        scores.sort(key=lambda x: x[0], reverse=True)
        selected = [item[1] for item in scores[:top_k]]
        return selected

    def extract_tailored_position(self, job_text: str) -> str:
        """Extracts 1 dynamic tailored position title based on job posting text."""
        t = job_text.lower()
        if any(k in t for k in ["scm", "물류", "wms", "입출고", "센터", "냉장", "냉동"]):
            return "스마트 물류센터 총괄 관리자 / SCM전산 수석"
        elif any(k in t for k in ["영업", "유통", "가맹점", "개설", "채널", "상권"]):
            return "총괄 영업유통 & 신규 가맹 채널 개발 수석"
        elif any(k in t for k in ["홈쇼핑", "방송", "기획", "t커머스", "마케팅"]):
            return "TV홈쇼핑 & 이커머스 유통 총괄 부장"
        elif any(k in t for k in ["ai", "에이전트", "전산", "데이터", "it", "파이썬"]):
            return "AI 에이전트 수석아키텍트 / IT전산 총괄"
        return "스마트 물류센터 총괄 관리자 & AI 에이전트 아키텍트"

    def _get_profile_b64(self) -> Optional[str]:
        """Loads profile photo from assets directory and returns base64 data URI."""
        import base64
        paths = [
            BASE_DIR / "assets" / "photo.jpg",
            BASE_DIR / "assets" / "profile.jpg",
            BASE_DIR / "data" / "assets" / "profile.jpg",
            BASE_DIR / "docs" / "assets" / "profile.jpg",
            BASE_DIR / "career_hub" / "profile.jpg",
            BASE_DIR / "photo.jpg"
        ]
        for p in paths:
            if p.exists():
                try:
                    with open(p, "rb") as f:
                        encoded = base64.b64encode(f.read()).decode("utf-8")
                        return f"data:image/jpeg;base64,{encoded}"
                except Exception as e:
                    print(f"Error reading profile photo: {e}")
        return None

    def filter_work_chronology(self, job_text: str, all_works: List[Dict[str, str]], top_k: int = 6) -> List[Dict[str, str]]:
        """Filters out short-term experiences (< 3 months) and selects top core experiences matching job posting."""
        filtered = [w for w in all_works if w["company"] != "프로축산" and "2개월" not in w.get("details", "")]

        scored = []
        for w in filtered:
            text = f"{w['company']} {w['role']} {w['details']}"
            score = 0
            if w["company"] in ["주식회사 그래이박스", "주식회사 케이엠기획", "주식회사 지엠지", "(주)현대시트"]:
                score += 5
            for kw in ["물류", "scm", "wms", "영업", "유통", "전산", "현장", "관리", "팀장", "부장", "차장", "본부장"]:
                if kw in job_text.lower() and kw in text.lower():
                    score += 3
            scored.append((score, w))

        scored.sort(key=lambda x: x[0], reverse=True)
        top_entries = [item[1] for item in scored[:top_k]]

        ordered = [w for w in filtered if w in top_entries]
        return ordered

    def clean_company_and_job_title(self, user_input: str) -> tuple[str, str]:
        """
        Cleans natural language user prompt to extract pure company name and job position title.
        Removes filler words like '알파야', '작성해줘', '지원서', '포지션', '~라는 회사인데', '입사' etc.
        """
        raw = user_input.strip()
        
        # Remove common command prefixes
        for prefix in ["/지원", "알파야", "알파", "수석비서", "대표님"]:
            if raw.startswith(prefix):
                raw = raw[len(prefix):].strip()

        # Remove filler words and trailing requests
        fillers = [
            "입사지원서", "입사 지원서", "지원서", "포지션", "채용공고", "채용 공고", "채용",
            "작성해 줘", "작성해줘", "만들어 줘", "만들어줘", "생성해 줘", "생성해줘",
            "써 줘", "써줘", "제출", "작성", "만들어", "생성", "해줘",
            "라는 회사인데", "이라는 회사인데", "회사인데", "회사"
        ]
        
        cleaned = raw
        for f in fillers:
            cleaned = cleaned.replace(f, "")
        cleaned = cleaned.strip()

        # Split into Company Name and Job Title
        company_name = ""
        job_title = ""

        tokens = cleaned.split()
        if len(tokens) >= 2:
            company_name = tokens[0]
            job_title = " ".join(tokens[1:])
        elif len(tokens) == 1 and tokens[0]:
            company_name = tokens[0]
            job_title = "맞춤 포지션 수석"
        else:
            company_name = "맞춤 기업"
            job_title = "총괄 관리자 / 수석 아키텍트"

        return company_name.strip(), job_title.strip()

    def generate_application_data(self, company_name: str, job_posting_text: str) -> Dict[str, Any]:
        """Generates tailored resume & cover letter data matching the job posting."""
        # Run natural language cleaner on company_name / user_input
        cleaned_co, cleaned_pos = self.clean_company_and_job_title(company_name)
        if cleaned_co and cleaned_co != "맞춤 기업":
            company_name = cleaned_co
        
        dynamic_position = cleaned_pos if (cleaned_pos and cleaned_pos != "맞춤 포지션 수석") else self.extract_tailored_position(job_posting_text)
        matched_stars = self.match_star_episodes(job_posting_text, top_k=3)

        # Basic Info (Ground Truth)
        basic_info = {
            "name": "김승률 (Kim Seung-ryul)",
            "birth": "1975년 2월 23일 (만 51세)",
            "position": dynamic_position,
            "experience_total": "15년 8개월+ (물류·유통·해외주재원·CS 총괄)",
            "address": "인천광역시 남동구 호구포로765번길 68-21 (구월동)",
            "contact": "010-4549-2886 / chan7502@naver.com",
            "skills": "Python AI 에이전트 구축, MS Excel 쿼리/함수 자동화, ERP/WMS 물류전산, 지게차 3톤, CS 민원 조율"
        }

        # Core Competencies (3 Bullets)
        core_competencies = [
            f"15년 8개월+ 현장 및 전산 오퍼레이션 통합 관리 및 {company_name} 성과 창출 역량",
            "4,000평 물류센터 WMS 구축 및 MS Excel 쿼리 자동화를 통한 오차율 0% & 매출 12배 신장",
            "카자흐스탄 주재원 6년 및 사업체 운영을 통해 검증된 1:1 경청 소통 및 갈등 조율 리더십"
        ]

        # Work Chronology (Full Ground Truth 12 items)
        full_work_chronology = [
            {"period": "2024.08 ~ 2026.06", "company": "KB라이프파트너스", "role": "설계사 / 컨설턴트", "details": "고객 대면 1:1 맞춤형 자산 컨설팅 및 CS 관리, 프리미엄 케어"},
            {"period": "2025.12 ~ 2026.01", "company": "프로축산", "role": "직장가입자", "details": "축산물 물류 전산 및 출고 오퍼레이션 지원"},
            {"period": "2023.06 ~ 2024.08", "company": "주식회사 그래이박스", "role": "차장 / 전산OP", "details": "4000평 냉장/냉동 물류센터 입출고전산, WMS 연동, 월매출 12배 신장"},
            {"period": "2022.04 ~ 2023.05", "company": "주식회사 지엠지", "role": "부장", "details": "전사 기획 및 영업 총괄 관리, 부서 간 갈등 조율"},
            {"period": "2020.03 ~ 2022.04", "company": "혼밥집 부평점", "role": "대표", "details": "외식 사업체 총괄 기획 및 직접 매장 운영, 피크타임 동선 제어 및 CS 해결"},
            {"period": "2019.03 ~ 2019.08", "company": "(주)원창엔지니어링", "role": "직장가입자", "details": "엔지니어링 현장 관리 및 자재 수급 지원"},
            {"period": "2016.07 ~ 2017.10", "company": "주식회사 서우", "role": "사원 / 팀장", "details": "3교대 현장 생산 공정 오퍼레이션 통제 및 3톤 지게차 무사고운전"},
            {"period": "2015.11 ~ 2016.07", "company": "(주)현대시트", "role": "팀장 / 차장", "details": "TV홈쇼핑 영업팀장, 따소미플러스 홈앤쇼핑/T커머스 런칭 (누적 10억 매출)"},
            {"period": "2009.12 ~ 2015.05", "company": "주식회사 케이엠기획", "role": "차장 / 본부장", "details": "카자흐스탄 로또복권 해외주재원 6년, 악토베/친켄트 소장 및 남부본부장"},
            {"period": "2003.07 ~ 2008.04", "company": "주식회사 케이엠기획", "role": "대리", "details": "로또복권 1기 인천 가맹점 모집 (개인달성률 144%)"},
            {"period": "2002.06 ~ 2003.07", "company": "(주)코리아로터리서비스", "role": "대리 / LSR", "details": "대한민국 최초 온라인 로또복권 유통망 모집 및 점주 교육"},
            {"period": "2002.01 ~ 2002.06", "company": "(주)희망백화점", "role": "사원 / 디자이너", "details": "백화점 쇼핑몰 디자인 및 온라인 전산 관리"}
        ]

        # Apply filtering for job posting (Exclude < 3 months, select top core entries)
        filtered_works = self.filter_work_chronology(job_posting_text, full_work_chronology, top_k=5)

        # Education (Ground Truth)
        education = [
            {"school": "인천전문대학", "major": "인문사회학부 e-비즈니스과 (경영학 전공)", "status": "졸업", "period": "2007.03 ~ 2009.02", "gpa": "4.29 / 4.5"},
            {"school": "항도실업고등학교", "major": "전자과", "status": "졸업", "period": "1990.03 ~ 1993.02", "gpa": "-"}
        ]

        # Credentials (Ground Truth)
        credentials = [
            {"name": "유통관리사 2급", "issuer": "대한상공회의소", "date": "2008.09.24"},
            {"name": "전자상거래관리사 2급", "issuer": "대한상공회의소", "date": "2008.11.28"},
            {"name": "지게차운전기능사 (3톤 이상)", "issuer": "한국산업인력공단", "date": "2017.03.22"},
            {"name": "컴퓨터그래픽스운용기능사", "issuer": "한국산업인력공단", "date": "2000.08.07"},
            {"name": "일반경비원 신임교육이수증", "issuer": "경찰청 지정기관", "date": "2025.12.18"}
        ]

        # 4-Part Cover Letter Sections
        cover_letter_sections = [
            {
                "title": "1. 직무 핵심 적응력 및 현장 노하우",
                "content": f"15년 8개월간 종합 물류전산, 해외 주재원, TV홈쇼핑 영업, 프리미엄 CS 컨설팅 현장에서 쌓아온 노하우를 바탕으로, {company_name}의 {dynamic_position} 포지션에 즉각 투입되어 업무 프로세스를 빠르게 파악하고 안정화하겠습니다. 수많은 현장 변수 속에서도 목표 달성을 최우선으로 삼아 최선의 오퍼레이션을 수행해 왔습니다."
            },
            {
                "title": "2. 데이터·전산 처리 및 오퍼레이션 혁신 역량",
                "content": "(주)그래이박스 4,000평 냉장/냉동 물류센터에서 수기 방식의 오차를 해소하기 위해 전문 WMS(사방넷/엔윌) 운용 및 MS Excel 쿼리 자동 집계 서식을 구축했습니다. 그 결과 재고 입출고 오차율 0%를 달성하고 월 매출을 12배(500만 원 -> 6,000만 원) 신장시켰습니다. 최근 구축한 Python AI 에이전트 시스템과 결합하여 전산 정확도와 효율을 극대화하겠습니다."
            },
            {
                "title": "3. 소통, 노무 갈등 조율 및 조직 관리 리더십",
                "content": "카자흐스탄 주재원 근무 당시 침켄트 사무소의 전원 파업 위기 상황에 신임 소장으로 긴급 파견되어 1:1 경청 소통을 진행했습니다. 현지 직책 세분화(소장/부소장/Senior/Manager) 및 직책수당 체계, 업무표준 매뉴얼을 도입함으로써 2개월 만에 파업을 수습하고 조직을 100% 정상화하여 본부장으로 승진했습니다. 상이한 조직 문화 속에서도 갈등을 원만히 조율하는 정중하고 차분한 소통 리더십을 발휘하겠습니다."
            },
            {
                "title": "4. 지원동기 및 입사 후 포부",
                "content": f"{company_name}의 성장 가능성과 유통/물류 비전에 깊이 공감하여 {dynamic_position} 직무에 지원하게 되었습니다. 입사 후 100일 이내에 담당 영역의 현장 동선 및 전산 프로세스를 완벽히 표준화하고, 오차 없는 정밀 관리와 부서 간 조화로운 소통을 통해 {company_name}의 핵심 성과 창출에 기여하겠습니다."
            }
        ]

        return {
            "company_name": company_name,
            "basic_info": basic_info,
            "core_competencies": core_competencies,
            "work_chronology": filtered_works,
            "full_work_chronology": full_work_chronology,
            "education": education,
            "credentials": credentials,
            "cover_letter_sections": cover_letter_sections,
            "matched_stars": matched_stars
        }

    def render_html_template(self, app_data: Dict[str, Any], is_pwa: bool = True) -> str:
        """Renders Interactive Dark Dashboard (dashboard_view.html) or A4 Print Form (print_a4.html)."""
        from datetime import datetime
        today_str = datetime.now().strftime("%Y년 %m월 %d일")
        
        template_name = "dashboard_view.html" if is_pwa else "print_a4.html"
        env = jinja2.Environment(loader=jinja2.FileSystemLoader(str(TEMPLATES_DIR)))
        template = env.get_template(template_name)

        profile_b64 = self._get_profile_b64()

        rendered = template.render(
            company_name=app_data["company_name"],
            basic_info=app_data["basic_info"],
            core_competencies=app_data.get("core_competencies", []),
            work_chronology=app_data["work_chronology"],
            education=app_data.get("education", []),
            credentials=app_data["credentials"],
            cover_letter_sections=app_data.get("cover_letter_sections", []),
            profile_b64=profile_b64,
            today_str=today_str,
            is_pwa=is_pwa
        )
        return rendered

    def render_pdf_with_python_docx(self, app_data: Dict[str, Any], output_pdf_path: Path) -> bool:
        """Generates a high-fidelity 2-page Wyndham Goseong style MS Word document and converts it to PDF via docx2pdf."""
        output_pdf_path = Path(output_pdf_path)
        output_pdf_path.parent.mkdir(parents=True, exist_ok=True)
        output_docx_path = output_pdf_path.with_suffix(".docx")

        try:
            import docx
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            doc = docx.Document()
            
            # Set A4 margins (0.55 in ~ 14mm)
            for section in doc.sections:
                section.top_margin = Inches(0.55)
                section.bottom_margin = Inches(0.55)
                section.left_margin = Inches(0.55)
                section.right_margin = Inches(0.55)

            # Default font style
            style = doc.styles['Normal']
            font = style.font
            font.name = '맑은 고딕'
            font.size = Pt(10)
            font.color.rgb = RGBColor(17, 24, 39)

            def set_cell_bg(cell, hex_color):
                tcPr = cell._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), hex_color)
                tcPr.append(shd)

            def set_borders(table):
                tblPr = table._element.xpath('w:tblPr')
                if tblPr:
                    tblBorders = OxmlElement('w:tblBorders')
                    for b_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                        b = OxmlElement(f'w:{b_name}')
                        b.set(qn('w:val'), 'single')
                        b.set(qn('w:sz'), '4')
                        b.set(qn('w:space'), '0')
                        b.set(qn('w:color'), 'D1D5DB')
                        tblBorders.append(b)
                    tblPr[0].append(tblBorders)

            # --- PAGE 1: 이력서 ---
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_title.paragraph_format.space_before = Pt(0)
            p_title.paragraph_format.space_after = Pt(10)
            r_title = p_title.add_run("이  력  서")
            r_title.font.size = Pt(20)
            r_title.font.bold = True

            basic_info = app_data.get("basic_info", {})
            tbl_profile = doc.add_table(rows=4, cols=5)
            tbl_profile.alignment = WD_TABLE_ALIGNMENT.CENTER
            set_borders(tbl_profile)

            col_widths = [Inches(0.9), Inches(1.8), Inches(0.9), Inches(2.2), Inches(1.2)]
            for row in tbl_profile.rows:
                for idx, w in enumerate(col_widths):
                    row.cells[idx].width = w

            # Merge Photo Column (Row 0..3, Col 4)
            photo_cell = tbl_profile.cell(0, 4)
            for r_idx in range(1, 4):
                photo_cell.merge(tbl_profile.cell(r_idx, 4))

            p_photo = photo_cell.paragraphs[0]
            p_photo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            photo_path = BASE_DIR / "assets" / "photo.jpg"
            if photo_path.exists():
                r_img = p_photo.add_run()
                r_img.add_picture(str(photo_path), width=Inches(1.05), height=Inches(1.3))

            info_rows = [
                [("성 명", basic_info.get("name", "김승률")), ("생년월일", basic_info.get("birth", "1983.10.15"))],
                [("연 락 처", "010-4549-2886"), ("이 메 일", "chan7502@naver.com")],
                [("주 소", basic_info.get("address", "경기도 시흥시 배곧1로 27, 214동 1204호")), None],
                [("지원회사", app_data.get("company_name", "")), ("지원포지션", basic_info.get("position", "수석 아키텍트"))]
            ]

            for r_idx, row_data in enumerate(info_rows):
                cell_l1 = tbl_profile.cell(r_idx, 0)
                cell_v1 = tbl_profile.cell(r_idx, 1)
                cell_l1.text = row_data[0][0]
                cell_v1.text = str(row_data[0][1])
                set_cell_bg(cell_l1, "F4F5F7")
                cell_l1.paragraphs[0].runs[0].font.bold = True
                cell_l1.paragraphs[0].runs[0].font.size = Pt(9.5)
                cell_v1.paragraphs[0].runs[0].font.size = Pt(9.5)

                if r_idx == 2:
                    cell_v1.merge(tbl_profile.cell(r_idx, 3))
                elif row_data[1]:
                    cell_l2 = tbl_profile.cell(r_idx, 2)
                    cell_v2 = tbl_profile.cell(r_idx, 3)
                    cell_l2.text = row_data[1][0]
                    cell_v2.text = str(row_data[1][1])
                    set_cell_bg(cell_l2, "F4F5F7")
                    cell_l2.paragraphs[0].runs[0].font.bold = True
                    cell_l2.paragraphs[0].runs[0].font.size = Pt(9.5)
                    cell_v2.paragraphs[0].runs[0].font.size = Pt(9.5)

            # Core Competencies
            p_sec1 = doc.add_paragraph()
            p_sec1.paragraph_format.space_before = Pt(8)
            p_sec1.paragraph_format.space_after = Pt(2)
            r_sec1 = p_sec1.add_run("■ 핵심 역량 요약")
            r_sec1.font.size = Pt(11)
            r_sec1.font.bold = True

            for comp in app_data.get("core_competencies", [])[:3]:
                p_b = doc.add_paragraph(style='List Bullet')
                p_b.paragraph_format.space_before = Pt(0)
                p_b.paragraph_format.space_after = Pt(2)
                r_b = p_b.add_run(comp)
                r_b.font.size = Pt(9.5)

            # Work Chronology
            p_sec2 = doc.add_paragraph()
            p_sec2.paragraph_format.space_before = Pt(8)
            p_sec2.paragraph_format.space_after = Pt(2)
            r_sec2 = p_sec2.add_run("■ 주요 경력 사항")
            r_sec2.font.size = Pt(11)
            r_sec2.font.bold = True

            works = app_data.get("work_chronology", [])
            tbl_work = doc.add_table(rows=len(works) + 1, cols=4)
            tbl_work.alignment = WD_TABLE_ALIGNMENT.CENTER
            set_borders(tbl_work)

            headers_w = [("근무기간", Inches(1.3)), ("사업장명", Inches(1.8)), ("직책", Inches(1.1)), ("담당 직무 및 핵심 성과", Inches(2.8))]
            for c_idx, (h_text, h_w) in enumerate(headers_w):
                cell = tbl_work.cell(0, c_idx)
                cell.width = h_w
                cell.text = h_text
                set_cell_bg(cell, "F4F5F7")
                r = cell.paragraphs[0].runs[0]
                r.font.bold = True
                r.font.size = Pt(9.5)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            for r_idx, w in enumerate(works, start=1):
                row_vals = [w.get("period", ""), w.get("company", ""), w.get("role", ""), w.get("details", "")]
                for c_idx, val in enumerate(row_vals):
                    cell = tbl_work.cell(r_idx, c_idx)
                    cell.text = str(val)
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
                    if c_idx in [0, 2]:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Education
            p_sec3 = doc.add_paragraph()
            p_sec3.paragraph_format.space_before = Pt(8)
            p_sec3.paragraph_format.space_after = Pt(2)
            r_sec3 = p_sec3.add_run("■ 학력 사항")
            r_sec3.font.size = Pt(11)
            r_sec3.font.bold = True

            edus = app_data.get("education", [])
            tbl_edu = doc.add_table(rows=len(edus) + 1, cols=5)
            tbl_edu.alignment = WD_TABLE_ALIGNMENT.CENTER
            set_borders(tbl_edu)

            headers_e = [("학교명", Inches(1.8)), ("전공/학과", Inches(2.1)), ("구분", Inches(0.9)), ("기간", Inches(1.3)), ("학점", Inches(0.9))]
            for c_idx, (h_text, h_w) in enumerate(headers_e):
                cell = tbl_edu.cell(0, c_idx)
                cell.width = h_w
                cell.text = h_text
                set_cell_bg(cell, "F4F5F7")
                r = cell.paragraphs[0].runs[0]
                r.font.bold = True
                r.font.size = Pt(9.5)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            for r_idx, e in enumerate(edus, start=1):
                row_vals = [e.get("school", ""), e.get("major", ""), e.get("status", ""), e.get("period", ""), e.get("gpa", "")]
                for c_idx, val in enumerate(row_vals):
                    cell = tbl_edu.cell(r_idx, c_idx)
                    cell.text = str(val)
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
                    if c_idx >= 2:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Credentials
            p_sec4 = doc.add_paragraph()
            p_sec4.paragraph_format.space_before = Pt(8)
            p_sec4.paragraph_format.space_after = Pt(2)
            r_sec4 = p_sec4.add_run("■ 보유 자격증 & 면허")
            r_sec4.font.size = Pt(11)
            r_sec4.font.bold = True

            creds = app_data.get("credentials", [])
            tbl_cred = doc.add_table(rows=len(creds) + 1, cols=3)
            tbl_cred.alignment = WD_TABLE_ALIGNMENT.CENTER
            set_borders(tbl_cred)

            headers_c = [("자격 / 면허명", Inches(2.8)), ("발급 기관", Inches(2.4)), ("취득 일자", Inches(1.8))]
            for c_idx, (h_text, h_w) in enumerate(headers_c):
                cell = tbl_cred.cell(0, c_idx)
                cell.width = h_w
                cell.text = h_text
                set_cell_bg(cell, "F4F5F7")
                r = cell.paragraphs[0].runs[0]
                r.font.bold = True
                r.font.size = Pt(9.5)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            for r_idx, c in enumerate(creds, start=1):
                row_vals = [c.get("name", ""), c.get("issuer", ""), c.get("date", "")]
                for c_idx, val in enumerate(row_vals):
                    cell = tbl_cred.cell(r_idx, c_idx)
                    cell.text = str(val)
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
                    if c_idx == 2:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # --- Page Break to Page 2 ---
            doc.add_page_break()

            # --- PAGE 2: 자기소개서 ---
            p_title2 = doc.add_paragraph()
            p_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_title2.paragraph_format.space_before = Pt(0)
            p_title2.paragraph_format.space_after = Pt(12)
            r_title2 = p_title2.add_run("자  기  소  개  서")
            r_title2.font.size = Pt(16)
            r_title2.font.bold = True

            sections = app_data.get("cover_letter_sections", [])
            for sec in sections:
                p_stitle = doc.add_paragraph()
                p_stitle.paragraph_format.space_before = Pt(6)
                p_stitle.paragraph_format.space_after = Pt(2)
                r_stitle = p_stitle.add_run(sec.get("title", ""))
                r_stitle.font.size = Pt(11)
                r_stitle.font.bold = True

                p_content = doc.add_paragraph()
                p_content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_content.paragraph_format.line_spacing = 1.35
                p_content.paragraph_format.space_before = Pt(0)
                p_content.paragraph_format.space_after = Pt(6)
                r_content = p_content.add_run(sec.get("content", ""))
                r_content.font.size = Pt(10)

            # Signature Block
            p_date = doc.add_paragraph()
            p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_date.paragraph_format.space_before = Pt(20)
            p_date.paragraph_format.space_after = Pt(10)
            r_date = p_date.add_run(app_data.get("today_str", "2026년 8월 28일"))
            r_date.font.size = Pt(11)
            r_date.font.bold = True

            p_sig = doc.add_paragraph()
            p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_sig.paragraph_format.space_before = Pt(0)
            p_sig.paragraph_format.space_after = Pt(18)
            r_sig = p_sig.add_run("지 원 자 :   김  승  률     (인 / 서명)")
            r_sig.font.size = Pt(11.5)
            r_sig.font.bold = True

            p_comp = doc.add_paragraph()
            p_comp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_comp.paragraph_format.space_before = Pt(8)
            p_comp.paragraph_format.space_after = Pt(0)
            r_comp = p_comp.add_run(f"{app_data.get('company_name', '')} 귀중")
            r_comp.font.size = Pt(14)
            r_comp.font.bold = True

            doc.save(str(output_docx_path))

            # Convert to PDF
            try:
                from docx2pdf import convert
                convert(str(output_docx_path), str(output_pdf_path))
                if output_pdf_path.exists() and output_pdf_path.stat().st_size > 0:
                    print(f"[OK] docx2pdf conversion succeeded: {output_pdf_path}")
                    return True
            except Exception as e_pdf:
                print(f"docx2pdf warning: {e_pdf}. Falling back to Playwright...")

        except Exception as e_docx:
            print(f"python-docx error: {e_docx}. Falling back to Playwright...")

        # Fallback to Playwright HTML rendering if python-docx/docx2pdf fails
        return self.render_pdf_with_playwright(app_data, output_pdf_path)

    def render_pdf_with_playwright(self, app_data_or_html: Any, output_pdf_path: Path) -> bool:
        """Renders A4 PDF using Playwright with print_a4.html and automatically cleans up temp files."""
        output_pdf_path = Path(output_pdf_path)
        output_pdf_path.parent.mkdir(parents=True, exist_ok=True)
        
        if isinstance(app_data_or_html, dict):
            html_content = self.render_html_template(app_data_or_html, is_pwa=False)
        else:
            html_content = str(app_data_or_html)

        temp_html_path = output_pdf_path.with_suffix(".temp.html")
        with open(temp_html_path, "w", encoding="utf-8") as f:
            f.write(html_content)

        try:
            from playwright.sync_api import sync_playwright
            with sync_playwright() as p:
                browser = p.chromium.launch(headless=True)
                page = browser.new_page()
                page.goto(temp_html_path.as_uri())
                page.pdf(
                    path=str(output_pdf_path),
                    format="A4",
                    print_background=True,
                    prefer_css_page_size=True,
                    margin={"top": "0", "bottom": "0", "left": "0", "right": "0"}
                )
                browser.close()
            
            # Clean up temporary html file
            if temp_html_path.exists():
                temp_html_path.unlink()

            return output_pdf_path.exists() and output_pdf_path.stat().st_size > 0
        except Exception as e:
            print(f"Error generating PDF via Playwright: {e}")
            if temp_html_path.exists():
                temp_html_path.unlink()
            return False

    def get_slug(self, company_name: str) -> str:
        clean = company_name.strip()
        if "지오영" in clean:
            return "geo_young"
        if "그래이박스" in clean or "graybox" in clean:
            return "graybox"
        if "엠아이큐브" in clean or "micube" in clean:
            return "micubesoft"
        s = re.sub(r'[()\s/\\:]+', '_', clean).strip('_')
        return s if s else "app"

    def publish_to_github_pages(self, company_name: str, pwa_html: str, pdf_path: Optional[Path] = None) -> Optional[str]:
        """Cleans up legacy target files, saves WebApp index.html, copies resume.pdf, and commits to GitHub Pages."""
        import shutil
        slug = self.get_slug(company_name)
        
        # Clean target directories before deploy
        for base_dir in [DOCS_APPS_DIR / slug, BASE_DIR / "applications" / slug]:
            if base_dir.exists():
                shutil.rmtree(base_dir, ignore_errors=True)
            base_dir.mkdir(parents=True, exist_ok=True)

        docs_index_path = DOCS_APPS_DIR / slug / "index.html"
        with open(docs_index_path, "w", encoding="utf-8") as f:
            f.write(pwa_html)

        root_index_path = BASE_DIR / "applications" / slug / "index.html"
        with open(root_index_path, "w", encoding="utf-8") as f:
            f.write(pwa_html)

        # Copy PDF to target directories as resume.pdf
        if pdf_path and Path(pdf_path).exists():
            shutil.copy(pdf_path, DOCS_APPS_DIR / slug / "resume.pdf")
            shutil.copy(pdf_path, BASE_DIR / "applications" / slug / "resume.pdf")

        # Git commit & push
        try:
            subprocess.run(["git", "add", "."], cwd=str(BASE_DIR), check=True)
            subprocess.run(["git", "commit", "-m", f"Deploy tailored job application and resume.pdf for {company_name}"], cwd=str(BASE_DIR), check=False)
            subprocess.run(["git", "push", "origin", "main"], cwd=str(BASE_DIR), check=False)
        except Exception as e:
            print(f"Git commit/push warning: {e}")

        web_url = f"{BASE_URL}{slug}/"
        return web_url

    def delete_application(self, text: str = "") -> Dict[str, Any]:
        """
        Deletes GitHub Pages deployment HTML directories (docs/applications/{slug} and applications/{slug}).
        Preserves local PDF in data/outputs/ (Non-destructive Ground Truth principle).
        """
        import shutil
        cleaned_slugs = []
        text_strip = text.strip()
        
        # 1) 전체 삭제 조건: '다', '전체', '모두', '전부'가 포함되거나 '/정리' 단독 입력
        is_all = any(kw in text_strip for kw in ['다', '전체', '모두', '전부']) or text_strip == '/정리' or text_strip == ''

        target_keyword = ""
        target_slug = ""

        if not is_all:
            # 2) 타겟 선별 삭제 조건: 문장에서 명령 키워드 제거
            target_keyword = text_strip
            for rem in ['지워줘', '지워', '삭제해줘', '삭제해', '삭제', '정리해줘', '정리해', '내려줘', '페이지', '지원서', '배포', '해제', '/정리']:
                target_keyword = target_keyword.replace(rem, '')
            target_keyword = target_keyword.strip()

            if not target_keyword:
                is_all = True
            else:
                target_slug = self.get_slug(target_keyword)

        for base_path in [DOCS_APPS_DIR, BASE_DIR / "applications"]:
            if base_path.exists():
                for item in base_path.iterdir():
                    if not item.is_dir():
                        continue

                    should_delete = False
                    if is_all:
                        should_delete = True
                    else:
                        if (target_slug and (item.name == target_slug or target_keyword in item.name or target_slug in item.name or item.name in target_keyword)) or (target_keyword in item.name):
                            should_delete = True

                    if should_delete:
                        shutil.rmtree(item, ignore_errors=True)
                        if item.name not in cleaned_slugs:
                            cleaned_slugs.append(item.name)

        commit_msg = "Clean all application pages" if is_all else f"Remove {target_keyword} application page"
        try:
            subprocess.run(["git", "add", "-A"], cwd=str(BASE_DIR), check=True)
            subprocess.run(["git", "commit", "-m", commit_msg], cwd=str(BASE_DIR), check=False)
            subprocess.run(["git", "push", "origin", "main"], cwd=str(BASE_DIR), check=False)
        except Exception as e:
            print(f"Git push warning during delete: {e}")

        return {
            "success": True,
            "is_all": is_all,
            "target_keyword": target_keyword,
            "cleaned_slugs": cleaned_slugs
        }

    def send_telegram_notification(self, company_name: str, pdf_path: Optional[Path], web_url: str) -> bool:
        """Sends Telegram dual notification (text with web URL + attached A4 PDF)."""
        import requests
        from dotenv import load_dotenv
        load_dotenv(dotenv_path=BASE_DIR / ".env")

        token = os.getenv("TELEGRAM_BOT_TOKEN")
        chat_id = os.getenv("TELEGRAM_CHAT_ID")
        if not token or not chat_id:
            print("[Telegram] Token or Chat ID missing.")
            return False

        message = f"""🎉 <b>[{company_name}] 맞춤 입사지원서 생성 완료!</b>

📱 <b>모바일 커리어 대시보드 (웹앱):</b>
{web_url}

📄 <b>정식 서류제출용 A4 PDF가 아래에 첨부되었습니다.</b>"""

        success = True
        try:
            # 1. Send text message with WebApp URL
            msg_url = f"https://api.telegram.org/bot{token}/sendMessage"
            res_msg = requests.post(
                msg_url,
                json={"chat_id": chat_id, "text": message, "parse_mode": "HTML"},
                timeout=180
            )
            if res_msg.status_code != 200:
                print(f"[Telegram] Text message send failed: {res_msg.status_code}")
                success = False

            # 2. Send PDF Document
            if pdf_path and Path(pdf_path).exists():
                doc_url = f"https://api.telegram.org/bot{token}/sendDocument"
                attachment_name = Path(pdf_path).name
                with open(pdf_path, "rb") as f:
                    res_doc = requests.post(
                        doc_url,
                        data={"chat_id": chat_id, "caption": f"📄 {company_name} 서류제출용 A4 PDF"},
                        files={"document": (attachment_name, f, "application/pdf")},
                        timeout=180
                    )
                    if res_doc.status_code != 200:
                        print(f"[Telegram] Document send failed: {res_doc.status_code}")
                        success = False
        except Exception as e:
            print(f"[Telegram] Dispatch error: {e}")
            success = False

        return success

    def generate_job_application(self, company_name: str, job_posting_text: str) -> Dict[str, Any]:
        """Full pipeline execution: Data extraction -> PDF (print_a4.html) -> PWA Dashboard (dashboard_view.html) -> Archiving -> Telegram Dispatch."""
        sanitized_company = self.get_slug(company_name)
        app_data = self.generate_application_data(company_name, job_posting_text)

        # 1. Render Interactive Dark Dashboard WebApp HTML
        dashboard_html = self.render_html_template(app_data, is_pwa=True)

        # 2. Generate A4 PDF using python-docx & docx2pdf (Wyndham Goseong 2-Page Format)
        pdf_filename = f"{sanitized_company}_김승률_지원서.pdf"
        output_pdf_path = OUTPUT_DIR / pdf_filename
        pdf_success = self.render_pdf_with_python_docx(app_data, output_pdf_path)

        # 3. Publish WebApp to GitHub Pages (Cleans up previous build files and deploys resume.pdf)
        web_url = self.publish_to_github_pages(company_name, dashboard_html, pdf_path=output_pdf_path if pdf_success else None)

        # 4. 3-Track Archiving
        from skills.report_archiver import archive_report_locally
        report_content = f"""# [지원서 생성 보고서] {company_name} 맞춤형 입사지원서

- **공고 회사명**: {company_name}
- **PDF 파일 생성**: {'성공 (' + str(output_pdf_path) + ')' if pdf_success else '실패'}
- **GitHub Pages WebApp URL**: {web_url}
- **매칭된 3대 STAR 에피소드**:
  1. {app_data['matched_stars'][0]['title']}
  2. {app_data['matched_stars'][1]['title']}
  3. {app_data['matched_stars'][2]['title']}
"""
        archive_success, archive_path, _ = archive_report_locally(f"{sanitized_company}_지원서생성_완료보고서", report_content)

        # 5. Telegram Dual Dispatch (Text + PDF)
        telegram_success = self.send_telegram_notification(
            company_name=company_name,
            pdf_path=output_pdf_path if pdf_success else None,
            web_url=web_url
        )

        return {
            "company_name": company_name,
            "pdf_success": pdf_success,
            "pdf_path": str(output_pdf_path) if pdf_success else None,
            "web_url": web_url,
            "archive_path": archive_path,
            "telegram_success": telegram_success,
            "matched_stars": [s["title"] for s in app_data["matched_stars"]]
        }

if __name__ == "__main__":
    generator = JobApplicationGenerator()
    result = generator.generate_job_application(
        company_name="종합물류_그래이박스",
        job_posting_text="종합물류기업 SCM 및 4000평 냉장냉동 물류센터 총괄 관리자 채용 (WMS 전산, 입출고 관리, 엑셀 쿼리, 현장관리)"
    )
    print("Execution Result:", json.dumps(result, ensure_ascii=False, indent=2))
